from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document as DocxDocument

from yiyan_dingzhen.documents import DocumentReadError, extract_text_from_bytes


def test_extract_plain_text() -> None:
    assert (
        extract_text_from_bytes(
            "物理文本".encode(),
            filename="sample.txt",
            content_type="text/plain",
        )
        == "物理文本"
    )


def test_extract_docx_paragraphs_and_tables() -> None:
    document = DocxDocument()
    document.add_paragraph("第一段")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "左"
    table.cell(0, 1).text = "右"
    buffer = BytesIO()
    document.save(buffer)

    text = extract_text_from_bytes(buffer.getvalue(), filename="sample.docx")
    assert "第一段" in text
    assert "左 右" in text


def test_reject_unsupported_document() -> None:
    with pytest.raises(DocumentReadError):
        extract_text_from_bytes(b"\x00\x01", filename="sample.bin")
