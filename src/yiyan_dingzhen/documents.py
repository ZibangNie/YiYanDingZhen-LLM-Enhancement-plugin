"""从受控字节流或本地路径提取文本。"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile


class DocumentReadError(ValueError):
    """文档格式不支持、损坏或超过安全限制。"""


_PDF_TYPES = {"application/pdf"}
_DOCX_TYPES = {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
_TEXT_TYPES = {
    "text/plain",
    "text/markdown",
    "text/x-markdown",
}


def get_file_extension_without_dot(file_path: str | Path) -> str:
    return Path(file_path).suffix.lower().lstrip(".")


def _bounded_text(parts: list[str], *, max_chars: int) -> str:
    text = "\n".join(part for part in parts if part)
    if len(text) > max_chars:
        raise DocumentReadError(f"提取后的文本超过 {max_chars} 个字符")
    return text


def _extract_pdf(data: bytes, *, max_pages: int, max_chars: int) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentReadError("缺少 pypdf 依赖") from exc

    try:
        reader = PdfReader(BytesIO(data), strict=False)
    except Exception as exc:
        raise DocumentReadError("PDF 无法解析") from exc
    if len(reader.pages) > max_pages:
        raise DocumentReadError(f"PDF 页数超过 {max_pages}")

    parts: list[str] = []
    total = 0
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            raise DocumentReadError("PDF 文本提取失败") from exc
        total += len(text)
        if total > max_chars:
            raise DocumentReadError(f"提取后的文本超过 {max_chars} 个字符")
        parts.append(text)
    return "\n".join(parts)


def _validate_docx_archive(data: bytes, *, max_uncompressed_bytes: int) -> None:
    try:
        with ZipFile(BytesIO(data)) as archive:
            total = sum(item.file_size for item in archive.infolist())
            if total > max_uncompressed_bytes:
                raise DocumentReadError(f"DOCX 解压后大小超过 {max_uncompressed_bytes} 字节")
    except BadZipFile as exc:
        raise DocumentReadError("DOCX 不是有效的 ZIP 文档") from exc


def _extract_docx(
    data: bytes,
    *,
    max_chars: int,
    max_uncompressed_bytes: int,
) -> str:
    _validate_docx_archive(data, max_uncompressed_bytes=max_uncompressed_bytes)
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise DocumentReadError("缺少 python-docx 依赖") from exc

    try:
        document = DocxDocument(BytesIO(data))
    except Exception as exc:
        raise DocumentReadError("DOCX 无法解析") from exc
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return _bounded_text(parts, max_chars=max_chars)


def extract_text_from_bytes(
    data: bytes,
    *,
    filename: str = "",
    content_type: str = "",
    max_pages: int = 500,
    max_chars: int = 2_000_000,
    max_docx_uncompressed_bytes: int = 100 * 1024 * 1024,
) -> str:
    if not data:
        raise DocumentReadError("文档内容为空")
    media_type = content_type.split(";", 1)[0].strip().lower()
    suffix = Path(filename).suffix.lower()

    if media_type in _PDF_TYPES or suffix == ".pdf" or data.startswith(b"%PDF"):
        return _extract_pdf(data, max_pages=max_pages, max_chars=max_chars)
    if media_type in _DOCX_TYPES or suffix == ".docx":
        return _extract_docx(
            data,
            max_chars=max_chars,
            max_uncompressed_bytes=max_docx_uncompressed_bytes,
        )
    if media_type in _TEXT_TYPES or suffix in {".txt", ".md", ".markdown"}:
        text = data.decode("utf-8", errors="replace")
        if len(text) > max_chars:
            raise DocumentReadError(f"文本超过 {max_chars} 个字符")
        return text
    raise DocumentReadError("仅支持 PDF、DOCX、Markdown 和纯文本文件")


def extract_text_from_path(
    path: str | Path,
    *,
    max_bytes: int = 20 * 1024 * 1024,
    max_chars: int = 2_000_000,
) -> str:
    resolved = Path(path).expanduser().resolve()
    if not resolved.is_file():
        raise DocumentReadError(f"文件不存在：{resolved}")
    size = resolved.stat().st_size
    if size > max_bytes:
        raise DocumentReadError(f"文件大小超过 {max_bytes} 字节")
    return extract_text_from_bytes(
        resolved.read_bytes(),
        filename=resolved.name,
        max_chars=max_chars,
    )


def select_file() -> str:
    import tkinter as tk
    from tkinter import filedialog

    root = tk.Tk()
    root.withdraw()
    try:
        return filedialog.askopenfilename(
            filetypes=[
                ("Supported documents", "*.pdf *.docx *.md *.txt"),
                ("PDF files", "*.pdf"),
                ("Word documents", "*.docx"),
                ("Text files", "*.md *.txt"),
            ]
        )
    finally:
        root.destroy()


def get_file_content(path: str | Path | None = None) -> str:
    selected = str(path) if path is not None else select_file()
    if not selected:
        return ""
    return extract_text_from_path(selected)
